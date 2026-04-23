# backend/main.py
"""Local web application entry point.

Run:
    python backend/main.py
"""

from __future__ import annotations

import cgi
import json
import logging
import threading
import time
import urllib.parse
import uuid
import webbrowser
from dataclasses import dataclass, field
from http import HTTPStatus
from http.server import BaseHTTPRequestHandler, ThreadingHTTPServer
from io import BytesIO
from pathlib import Path
from typing import cast

from docx import Document
from docx.shared import Pt

from generator import AIRequestError, generate_pages, regenerate_page
from logger import init_logger
from model import extract_text_from_upload
from schemas import GenerationResult, PageContent
from settings import UI_SETTINGS


_LOGGER = logging.getLogger(__name__)
_BASE_DIR = Path(__file__).resolve().parent.parent
_FRONTEND_DIR = _BASE_DIR / "frontend"
_FRONT_JS_PATH = _FRONTEND_DIR / "front.js"


@dataclass(slots=True)
class GenerationJob:
    """Background generation job state."""

    job_id: str
    description: str
    total_pages: int
    reference_mode: bool
    reference_text: str
    status: str = "queued"
    pages: list[PageContent] = field(default_factory=list)
    error: str = ""
    created_at: float = field(default_factory=time.time)


class AppState:
    """Thread-safe application state container."""

    def __init__(self) -> None:
        """Initialize app state."""
        self._lock = threading.Lock()
        self._jobs: dict[str, GenerationJob] = {}

    def create_job(
        self,
        description: str,
        total_pages: int,
        reference_mode: bool,
        reference_text: str,
    ) -> GenerationJob:
        """Create and store a new generation job."""
        job = GenerationJob(
            job_id=uuid.uuid4().hex,
            description=description,
            total_pages=total_pages,
            reference_mode=reference_mode,
            reference_text=reference_text,
        )
        with self._lock:
            self._jobs[job.job_id] = job
        return job

    def get_job(self, job_id: str) -> GenerationJob | None:
        """Get a job by id."""
        with self._lock:
            return self._jobs.get(job_id)

    def set_job_running(self, job_id: str) -> None:
        """Mark job as running."""
        with self._lock:
            job = self._jobs.get(job_id)
            if job:
                job.status = "running"

    def append_job_page(self, job_id: str, page: PageContent) -> None:
        """Append one generated page to a job.

        原理：后台线程逐页产出并立即写入共享状态，前端轮询后就能逐页看到结果。
        """
        with self._lock:
            job = self._jobs.get(job_id)
            if not job:
                return

            replaced = False
            for index, existing in enumerate(job.pages):
                if existing.index == page.index:
                    job.pages[index] = page
                    replaced = True
                    break

            if not replaced:
                job.pages.append(page)
                job.pages.sort(key=lambda item: item.index)

    def set_job_completed(self, job_id: str) -> None:
        """Mark job as completed."""
        with self._lock:
            job = self._jobs.get(job_id)
            if job:
                job.status = "completed"

    def set_job_failed(self, job_id: str, error: str) -> None:
        """Mark job as failed."""
        with self._lock:
            job = self._jobs.get(job_id)
            if job:
                job.status = "failed"
                job.error = error

    def snapshot_job(self, job_id: str) -> dict[str, object]:
        """Return safe snapshot of a job for JSON response."""
        with self._lock:
            job = self._jobs.get(job_id)
            if not job:
                raise ValueError("job_id 不存在。")

            pages = [page.to_dict() for page in sorted(job.pages, key=lambda item: item.index)]
            return {
                "job_id": job.job_id,
                "status": job.status,
                "total_pages": job.total_pages,
                "completed_pages": len(job.pages),
                "pages": pages,
                "error": job.error,
            }


class PBLHTTPServer(ThreadingHTTPServer):
    """HTTP server with injected app state."""

    def __init__(self, server_address: tuple[str, int], handler_cls: type[BaseHTTPRequestHandler]) -> None:
        """Initialize server and app state."""
        super().__init__(server_address, handler_cls)
        self.app_state = AppState()


def _run_generation_job(app_state: AppState, job_id: str) -> None:
    """Run one generation job in background."""
    job = app_state.get_job(job_id)
    if job is None:
        return

    app_state.set_job_running(job_id)
    try:
        for page_index in range(job.total_pages):
            page = regenerate_page(
                description=job.description,
                total_pages=job.total_pages,
                page_index=page_index,
                reference_mode=job.reference_mode,
                reference_text=job.reference_text,
            )
            app_state.append_job_page(job_id, page)

        app_state.set_job_completed(job_id)
    except Exception as exc:  # noqa: BLE001
        _LOGGER.exception("Background generation failed: %s", exc)
        app_state.set_job_failed(job_id, str(exc))


def _build_index_html() -> str:
    """Build HTML page at runtime to keep entry simple.

    原理：由后端直接输出页面骨架，减少独立 HTML 文件依赖，确保 main.py 即可启动 UI。
    """
    return """<!DOCTYPE html>
<html lang="zh-CN">
<head>
  <meta charset="utf-8" />
  <meta name="viewport" content="width=device-width,initial-scale=1" />
  <title>PBL 教学教案生成系统</title>
  <style>
    * { box-sizing: border-box; }
    body {
      margin: 0;
      background: #f6f8fb;
      color: #1f2937;
      font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue",
        Arial, "PingFang SC", "Microsoft YaHei", sans-serif;
    }
    .app {
      max-width: 1180px;
      margin: 0 auto;
      padding: 24px;
    }
    .header {
      margin-bottom: 20px;
    }
    .header h1 {
      margin: 0 0 8px;
      font-size: 30px;
    }
    .header p {
      margin: 0;
      color: #4b5563;
      line-height: 1.6;
    }
    .card {
      background: #ffffff;
      border-radius: 16px;
      box-shadow: 0 10px 30px rgba(31, 41, 55, 0.08);
      padding: 20px;
      margin-bottom: 20px;
    }
    .grid {
      display: grid;
      gap: 16px;
      grid-template-columns: 1fr 280px;
    }
    @media (max-width: 960px) {
      .grid {
        grid-template-columns: 1fr;
      }
    }
    .label {
      display: block;
      margin-bottom: 8px;
      font-weight: 700;
    }
    textarea, input[type="number"] {
      width: 100%;
      border: 1px solid #d1d5db;
      border-radius: 12px;
      padding: 12px 14px;
      font-size: 14px;
      line-height: 1.6;
      background: #fff;
      color: #111827;
      outline: none;
    }
    textarea:focus, input[type="number"]:focus {
      border-color: #3b82f6;
      box-shadow: 0 0 0 4px rgba(59,130,246,0.10);
    }
    textarea {
      min-height: 220px;
      resize: vertical;
    }
    .toolbar {
      display: flex;
      flex-wrap: wrap;
      gap: 12px;
      align-items: center;
      margin-top: 16px;
    }
    .btn {
      border: none;
      border-radius: 12px;
      padding: 10px 16px;
      font-size: 14px;
      font-weight: 700;
      cursor: pointer;
      transition: transform 0.06s ease, opacity 0.2s ease;
    }
    .btn:active { transform: translateY(1px); }
    .btn:disabled { opacity: 0.6; cursor: not-allowed; }
    .btn-primary { background: #2563eb; color: #fff; }
    .btn-secondary { background: #e5e7eb; color: #111827; }
    .btn-success { background: #059669; color: #fff; }
    .btn-warning { background: #d97706; color: #fff; }
    .muted {
      color: #6b7280;
      font-size: 13px;
    }
    .status {
      margin-top: 12px;
      padding: 10px 12px;
      border-radius: 10px;
      background: #eff6ff;
      color: #1d4ed8;
      display: none;
      white-space: pre-wrap;
    }
    .status.error {
      background: #fef2f2;
      color: #b91c1c;
    }
    .status.visible {
      display: block;
    }
    .pages {
      display: grid;
      gap: 16px;
    }
    .page-card {
      background: #ffffff;
      border-radius: 16px;
      box-shadow: 0 10px 30px rgba(31, 41, 55, 0.08);
      overflow: hidden;
      border: 1px solid #eef2f7;
    }
    .page-head {
      display: flex;
      gap: 12px;
      justify-content: space-between;
      align-items: center;
      padding: 16px 20px;
      border-bottom: 1px solid #eef2f7;
      background: #fbfdff;
    }
    .page-head h3 {
      margin: 0;
      font-size: 20px;
    }
    .page-actions {
      display: flex;
      flex-wrap: wrap;
      gap: 8px;
    }
    .page-body {
      padding: 18px 20px 20px;
    }
    .page-textarea {
      width: 100%;
      min-height: 280px;
      border: 1px solid #dbe2ea;
      border-radius: 12px;
      padding: 14px;
      line-height: 1.7;
      resize: vertical;
      font-size: 14px;
    }
    .inline {
      display: flex;
      align-items: center;
      gap: 10px;
      flex-wrap: wrap;
    }
    .file-info {
      font-size: 13px;
      color: #4b5563;
    }
  </style>
</head>
<body>
  <div class="app">
    <div class="header">
      <h1>PBL 教学教案生成系统</h1>
      <p>运行 <code>python backend/main.py</code> 后，本地自动打开网页。支持逐页生成、单页重生成、上传 PDF/DOCX 作为参考文档，以及导出 DOCX。</p>
    </div>

    <div class="card">
      <div class="grid">
        <div>
          <label class="label" for="description">课程描述 / 生成需求</label>
          <textarea id="description">基于真实特殊教育场景，设计一个跨学科 PBL 教学教案生成系统。内容需要覆盖课程目标、活动设计、角色分工、阶段性产出物、评价方式与反思建议，适用于高校课程或教师培训场景。</textarea>
        </div>
        <div>
          <label class="label" for="totalPages">生成页数</label>
          <input id="totalPages" type="number" min="1" max="20" step="1" value="5" />
          <div class="muted" style="margin-top: 8px;">建议 3-10 页，生成过程中会逐页显示。</div>

          <div style="height: 16px;"></div>

          <label class="label" for="referenceFile">参考文档（可选）</label>
          <input id="referenceFile" type="file" accept=".pdf,.docx" />
          <div class="file-info" id="fileInfo">未上传参考文件。</div>

          <div style="height: 16px;"></div>

          <div class="inline">
            <input id="useReference" type="checkbox" />
            <label for="useReference">生成时使用参考文档风格</label>
          </div>

          <div class="muted" style="margin-top: 8px;">
            上传后会提取 PDF/DOCX 文本，用于生成“风格与结构类似”的教案。
          </div>
        </div>
      </div>

      <div class="toolbar">
        <button id="generateBtn" class="btn btn-primary" type="button">逐页生成教案</button>
        <button id="exportBtn" class="btn btn-success" type="button">导出 DOCX</button>
        <button id="clearBtn" class="btn btn-secondary" type="button">清空结果</button>
      </div>

      <div id="status" class="status"></div>
    </div>

    <div id="pages" class="pages"></div>
  </div>

  <script src="/static/front.js"></script>
</body>
</html>
"""


def _json_bytes(payload: dict[str, object]) -> bytes:
    """Encode JSON using UTF-8."""
    return json.dumps(payload, ensure_ascii=False).encode("utf-8")


def _read_request_json(handler: BaseHTTPRequestHandler) -> dict[str, object]:
    """Read JSON payload from current request."""
    length = int(handler.headers.get("Content-Length", "0"))
    body = handler.rfile.read(length)
    try:
        payload = json.loads(body.decode("utf-8"))
    except (UnicodeDecodeError, json.JSONDecodeError) as exc:
        raise ValueError("请求体不是合法 JSON。") from exc

    if not isinstance(payload, dict):
        raise ValueError("请求体必须是 JSON 对象。")
    return cast(dict[str, object], payload)


def _send_response(
    handler: BaseHTTPRequestHandler,
    status: HTTPStatus,
    content_type: str,
    body: bytes,
    extra_headers: dict[str, str] | None = None,
) -> None:
    """Send HTTP response with common headers."""
    handler.send_response(status.value)
    handler.send_header("Content-Type", content_type)
    handler.send_header("Content-Length", str(len(body)))
    handler.send_header("Cache-Control", "no-store")
    if extra_headers:
        for key, value in extra_headers.items():
            handler.send_header(key, value)
    handler.end_headers()
    handler.wfile.write(body)


def _send_json(
    handler: BaseHTTPRequestHandler,
    status: HTTPStatus,
    payload: dict[str, object],
) -> None:
    """Send JSON response."""
    _send_response(
        handler=handler,
        status=status,
        content_type="application/json; charset=utf-8",
        body=_json_bytes(payload),
    )


def _sanitize_pages(raw_pages: object) -> list[PageContent]:
    """Validate and convert JSON pages into typed pages."""
    if not isinstance(raw_pages, list):
        raise ValueError("pages 必须是数组。")

    pages: list[PageContent] = []
    for item in raw_pages:
        if not isinstance(item, dict):
            raise ValueError("pages 中的元素必须是对象。")
        index = item.get("index")
        title = item.get("title")
        body = item.get("body")
        if not isinstance(index, int):
            raise ValueError("page.index 必须是整数。")
        if not isinstance(title, str) or not isinstance(body, str):
            raise ValueError("page.title 和 page.body 必须是字符串。")
        pages.append(PageContent(index=index, title=title, body=body))
    return pages


def _build_docx_bytes(pages: list[PageContent], document_title: str) -> bytes:
    """Build DOCX bytes from pages."""
    document = Document()
    normal_style = document.styles["Normal"]
    normal_style.font.name = "Arial"
    normal_style.font.size = Pt(11)

    title_paragraph = document.add_paragraph()
    title_run = title_paragraph.add_run(document_title)
    title_run.bold = True
    title_run.font.size = Pt(18)

    document.add_paragraph("")

    for page in pages:
        heading = document.add_paragraph()
        heading_run = heading.add_run(page.title)
        heading_run.bold = True
        heading_run.font.size = Pt(14)

        body_paragraph = document.add_paragraph()
        body_paragraph.add_run(page.body)
        document.add_paragraph("")

    output = BytesIO()
    document.save(output)
    return output.getvalue()


class PBLRequestHandler(BaseHTTPRequestHandler):
    """HTTP request handler for the local web app."""

    server_version = "PBLGenerationHTTP/1.1"

    @property
    def app_state(self) -> AppState:
        """Return typed app state."""
        return cast(PBLHTTPServer, self.server).app_state

    def do_GET(self) -> None:
        """Handle GET requests."""
        parsed = urllib.parse.urlparse(self.path)
        path = parsed.path

        if path == "/":
            body = _build_index_html().encode("utf-8")
            _send_response(
                handler=self,
                status=HTTPStatus.OK,
                content_type="text/html; charset=utf-8",
                body=body,
            )
            return

        if path == "/static/front.js":
            if not _FRONT_JS_PATH.exists():
                _send_json(
                    self,
                    HTTPStatus.NOT_FOUND,
                    {"error": "front.js 文件不存在。"},
                )
                return

            body = _FRONT_JS_PATH.read_bytes()
            _send_response(
                handler=self,
                status=HTTPStatus.OK,
                content_type="application/javascript; charset=utf-8",
                body=body,
            )
            return

        if path == "/api/job-status":
            self._handle_job_status(parsed.query)
            return

        _send_json(self, HTTPStatus.NOT_FOUND, {"error": "Not Found"})

    def do_POST(self) -> None:
        """Handle POST requests."""
        parsed = urllib.parse.urlparse(self.path)
        path = parsed.path

        try:
            if path == "/api/generate":
                self._handle_generate()
                return
            if path == "/api/generate-async":
                self._handle_generate_async()
                return
            if path == "/api/regenerate":
                self._handle_regenerate()
                return
            if path == "/api/export-docx":
                self._handle_export_docx()
                return
            if path == "/api/upload-source":
                self._handle_upload_source()
                return
            _send_json(self, HTTPStatus.NOT_FOUND, {"error": "Not Found"})
        except ValueError as exc:
            _send_json(self, HTTPStatus.BAD_REQUEST, {"error": str(exc)})
        except AIRequestError as exc:
            _send_json(self, HTTPStatus.BAD_GATEWAY, {"error": str(exc)})
        except Exception as exc:  # noqa: BLE001
            _LOGGER.exception("Unhandled server error: %s", exc)
            _send_json(
                self,
                HTTPStatus.INTERNAL_SERVER_ERROR,
                {"error": f"服务器内部错误：{exc}"},
            )

    def _handle_generate(self) -> None:
        """Handle full synchronous generation request."""
        payload = _read_request_json(self)

        description = payload.get("description")
        total_pages = payload.get("total_pages")
        reference_mode = payload.get("reference_mode", False)
        reference_text = payload.get("reference_text", "")

        if not isinstance(description, str) or not description.strip():
            raise ValueError("description 不能为空。")
        if not isinstance(total_pages, int):
            raise ValueError("total_pages 必须是整数。")
        if not isinstance(reference_mode, bool):
            raise ValueError("reference_mode 必须是布尔值。")
        if not isinstance(reference_text, str):
            raise ValueError("reference_text 必须是字符串。")

        pages = generate_pages(
            description=description,
            total_pages=total_pages,
            reference_mode=reference_mode,
            reference_text=reference_text,
        )
        result = GenerationResult(pages=pages)
        _send_json(self, HTTPStatus.OK, result.to_dict())

    def _handle_generate_async(self) -> None:
        """Handle async generation request."""
        payload = _read_request_json(self)

        description = payload.get("description")
        total_pages = payload.get("total_pages")
        reference_mode = payload.get("reference_mode", False)
        reference_text = payload.get("reference_text", "")

        if not isinstance(description, str) or not description.strip():
            raise ValueError("description 不能为空。")
        if not isinstance(total_pages, int) or total_pages <= 0:
            raise ValueError("total_pages 必须是正整数。")
        if not isinstance(reference_mode, bool):
            raise ValueError("reference_mode 必须是布尔值。")
        if not isinstance(reference_text, str):
            raise ValueError("reference_text 必须是字符串。")

        job = self.app_state.create_job(
            description=description,
            total_pages=total_pages,
            reference_mode=reference_mode,
            reference_text=reference_text,
        )

        worker = threading.Thread(
            target=_run_generation_job,
            args=(self.app_state, job.job_id),
            daemon=True,
        )
        worker.start()

        _send_json(
            self,
            HTTPStatus.OK,
            {
                "job_id": job.job_id,
                "status": job.status,
                "total_pages": job.total_pages,
            },
        )

    def _handle_job_status(self, query: str) -> None:
        """Handle job status query."""
        params = urllib.parse.parse_qs(query)
        job_id_list = params.get("job_id", [])
        if not job_id_list or not isinstance(job_id_list[0], str):
            raise ValueError("缺少 job_id。")

        snapshot = self.app_state.snapshot_job(job_id_list[0])
        _send_json(self, HTTPStatus.OK, snapshot)

    def _handle_regenerate(self) -> None:
        """Handle single page regeneration request."""
        payload = _read_request_json(self)

        description = payload.get("description")
        total_pages = payload.get("total_pages")
        page_index = payload.get("page_index")
        reference_mode = payload.get("reference_mode", False)
        reference_text = payload.get("reference_text", "")

        if not isinstance(description, str) or not description.strip():
            raise ValueError("description 不能为空。")
        if not isinstance(total_pages, int):
            raise ValueError("total_pages 必须是整数。")
        if not isinstance(page_index, int):
            raise ValueError("page_index 必须是整数。")
        if not isinstance(reference_mode, bool):
            raise ValueError("reference_mode 必须是布尔值。")
        if not isinstance(reference_text, str):
            raise ValueError("reference_text 必须是字符串。")

        page = regenerate_page(
            description=description,
            total_pages=total_pages,
            page_index=page_index,
            reference_mode=reference_mode,
            reference_text=reference_text,
        )
        _send_json(self, HTTPStatus.OK, page.to_dict())

    def _handle_export_docx(self) -> None:
        """Handle DOCX export request."""
        payload = _read_request_json(self)

        title = payload.get("title", "PBL 教学教案")
        pages_raw = payload.get("pages")

        if not isinstance(title, str) or not title.strip():
            raise ValueError("title 不能为空。")

        pages = _sanitize_pages(pages_raw)
        docx_bytes = _build_docx_bytes(pages=pages, document_title=title.strip())
        file_name = urllib.parse.quote(f"{title.strip()}.docx")

        _send_response(
            handler=self,
            status=HTTPStatus.OK,
            content_type=(
                "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
            ),
            body=docx_bytes,
            extra_headers={
                "Content-Disposition": f"attachment; filename*=UTF-8''{file_name}",
            },
        )

    def _handle_upload_source(self) -> None:
        """Handle source document upload request."""
        content_type = self.headers.get("Content-Type", "")
        if "multipart/form-data" not in content_type:
            raise ValueError("上传接口需要 multipart/form-data。")

        max_bytes = UI_SETTINGS.upload_max_size_mb * 1024 * 1024
        content_length = int(self.headers.get("Content-Length", "0"))
        if content_length <= 0:
            raise ValueError("未检测到上传内容。")
        if content_length > max_bytes:
            raise ValueError(
                f"上传文件过大，最大允许 {UI_SETTINGS.upload_max_size_mb}MB。"
            )

        environ = {
            "REQUEST_METHOD": "POST",
            "CONTENT_TYPE": content_type,
            "CONTENT_LENGTH": str(content_length),
        }
        form = cgi.FieldStorage(
            fp=self.rfile,
            headers=self.headers,
            environ=environ,
            keep_blank_values=False,
        )

        file_item = form["file"] if "file" in form else None
        if file_item is None or not getattr(file_item, "filename", ""):
            raise ValueError("请先选择要上传的 PDF 或 DOCX 文件。")

        file_name = str(file_item.filename)
        file_data = file_item.file.read()
        if not isinstance(file_data, bytes) or not file_data:
            raise ValueError("上传文件为空。")

        extracted_text = extract_text_from_upload(file_name=file_name, data=file_data)
        preview = extracted_text[:1000]

        _send_json(
            self,
            HTTPStatus.OK,
            {
                "file_name": file_name,
                "reference_text": extracted_text,
                "preview": preview,
                "message": "参考文件上传成功。",
            },
        )

    def log_message(self, format: str, *args: object) -> None:
        """Redirect HTTP logs to logging module."""
        _LOGGER.info("%s - %s", self.address_string(), format % args)


def _open_browser_later(url: str) -> None:
    """Open browser after server start."""
    if not UI_SETTINGS.auto_open_browser:
        return
    time.sleep(1.0)
    webbrowser.open(url)


def _start_server() -> None:
    """Start local HTTP server and open browser.

    原理：使用标准库 ThreadingHTTPServer 提供静态页面和 JSON API，
    并在独立线程里自动拉起浏览器，实现“运行 main.py 即打开界面”。
    """
    init_logger()
    server_address = (UI_SETTINGS.host, UI_SETTINGS.port)
    httpd = PBLHTTPServer(server_address, PBLRequestHandler)

    url = f"http://{UI_SETTINGS.host}:{UI_SETTINGS.port}"
    _LOGGER.info("PBL web app is running at %s", url)

    opener = threading.Thread(target=_open_browser_later, args=(url,), daemon=True)
    opener.start()

    try:
        httpd.serve_forever()
    except KeyboardInterrupt:
        _LOGGER.info("Received KeyboardInterrupt, shutting down.")
    finally:
        httpd.server_close()


if __name__ == "__main__":
    _start_server()